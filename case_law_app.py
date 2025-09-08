import streamlit as st
from fpdf import FPDF
from docx import Document
import pandas as pd
import io
from PIL import Image
import requests

# ---------------- PAGE CONFIG ----------------
st.set_page_config("Student Hub | LLB First Year", ":mortar_board:", layout="wide")

# ---------------- SESSION STATE ----------------
if "search_history" not in st.session_state:
    st.session_state.search_history = []
if "syllabus_files" not in st.session_state:
    st.session_state.syllabus_files = {}  # store uploaded syllabus files per subject
if "timetable_files" not in st.session_state:
    st.session_state.timetable_files = []  # store uploaded timetable/resources files
if "resource_links" not in st.session_state:
    st.session_state.resource_links = []  # store user input links

# ---------------- HEADER ----------------
def display_header():
    try:
        st.image(Image.open("logo.png"), width=100)
    except:
        st.markdown("""
        <div style='text-align:center;'>
            <h1 style='margin-bottom:0;'>🎓 Student Hub</h1>
            <p style='font-size:16px;color:#555;'>Syllabus, Timetable & Legal Acts</p>
        </div>""", unsafe_allow_html=True)
display_header()

# ---------------- EXPORT FUNCTIONS ----------------
def export_pdf(title: str, content_list: list) -> bytes:
    pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, title, ln=True, align="C"); pdf.ln(10)
    for i, item in enumerate(content_list, 1):
        if isinstance(item, dict):
            pdf.multi_cell(0, 10, f"{i}. {item.get('Title','')} - {item.get('Link','')}")
        else:
            pdf.multi_cell(0, 10, f"{i}. {item}")
    return pdf.output(dest="S").encode("latin1")

def export_docx(title: str, content_list: list) -> io.BytesIO:
    doc = Document(); doc.add_heading(title, 0)
    for i, item in enumerate(content_list, 1):
        if isinstance(item, dict):
            doc.add_paragraph(f"{i}. {item.get('Title','')}", style="List Number")
            doc.add_paragraph(f"Link: {item.get('Link','')}")
        else:
            doc.add_paragraph(f"{i}. {item}")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def export_excel(df: pd.DataFrame) -> io.BytesIO:
    buf = io.BytesIO(); df.to_excel(buf, index=False, sheet_name="Sheet1", engine="xlsxwriter"); buf.seek(0)
    return buf

# ---------------- UTILS ----------------
def fetch_pdf(url: str, act_name: str) -> bytes | None:
    try:
        return requests.get(url, timeout=15).content
    except:
        st.warning(
            f"⚠️ Unable to fetch '{act_name}' automatically. IndiaCode server may be down. "
            f"Check Google News or open manually 👉 [Link]({url})"
        )
        return None

# ---------------- TABS ----------------
tab1, tab2, tab3 = st.tabs(["📚 Syllabus & Notes", "🗓️ Timetable & Resources", "⚖️ Case Law / Act Search"])

# ---------------- TAB 1: SYLLABUS ----------------
with tab1:
    st.header("LLB(Hons) First Year Syllabus")

    subjects = [
        "Constitutional Law-I",
        "Environmental Law",
        "Law of Crimes-I (BNS)",
        "Family Law-I",
        "Law of Torts",
        "Contract-I"
    ]

    for subj in subjects:
        st.subheader(subj)
        uploaded_file = st.file_uploader(f"Upload PDF/DOCX for {subj}", type=["pdf","docx"], key=subj)
        if uploaded_file:
            st.session_state.syllabus_files[subj] = uploaded_file
            st.success(f"✅ Uploaded: {uploaded_file.name}")

        # Provide download button if file exists
        if subj in st.session_state.syllabus_files:
            file = st.session_state.syllabus_files[subj]
            st.download_button(
                f"📥 Download uploaded syllabus for {subj}", 
                data=file, 
                file_name=file.name,
                mime="application/pdf" if file.type=="application/pdf" 
                     else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ---------------- TAB 2: TIMETABLE & RESOURCES ----------------
with tab2:
    st.header("Class Timetable & Resources")
    
    # Upload files (Excel/CSV, PDF/DOCX, Image)
    uploaded_files = st.file_uploader(
        "Upload timetable or resources (Excel/CSV, PDF/DOCX, Image)", 
        type=["xlsx","csv","pdf","docx","jpg","jpeg","png"], 
        accept_multiple_files=True
    )
    if uploaded_files:
        st.session_state.timetable_files.extend(uploaded_files)
    
    # Display uploaded files
    for file in st.session_state.timetable_files:
        st.subheader(file.name)
        # Excel/CSV
        if file.type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet","text/csv"]:
            try:
                if file.name.endswith(".csv"):
                    df = pd.read_csv(file)
                else:
                    df = pd.read_excel(file)
                st.dataframe(df)
                st.download_button("📄 Export as PDF", export_pdf("Class Timetable", df.values.tolist()),
                                   file.name.replace(".csv",".pdf").replace(".xlsx",".pdf"), "application/pdf")
                st.download_button("📊 Export as Excel", export_excel(df),
                                   file.name.replace(".csv",".xlsx").replace(".xlsx",".xlsx"),
                                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            except Exception as e:
                st.error(f"Error reading file: {e}")
        # Images
        elif file.type.startswith("image"):
            st.image(file, width=400)
            st.download_button(f"📥 Download {file.name}", data=file, file_name=file.name, mime=file.type)
        # PDF/DOCX
        elif file.type in ["application/pdf","application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            st.download_button(f"📥 Download {file.name}", data=file, file_name=file.name, mime=file.type)
    
    # Option 1: Input links (YouTube/website)
    st.markdown("### Add Resource Links")
    link_input = st.text_input("Enter URL (YouTube, website, etc.)")
    if st.button("➕ Add Link"):
        if link_input:
            st.session_state.resource_links.append(link_input)
            st.success(f"Added link: {link_input}")
    
    # Display links
    if st.session_state.resource_links:
        st.markdown("### Your Resource Links")
        for i, link in enumerate(st.session_state.resource_links, 1):
            st.markdown(f"{i}. [🔗 {link}]({link})")

# ---------------- TAB 3: CASE LAW / ACT SEARCH ----------------
with tab3:
    st.header("Indian Legal Act Search Tool")

    known_acts = {
    "Constitution of India": "https://legislative.gov.in/sites/default/files/COI_2024.pdf",
    "Indian Penal Code 1860": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1860-45.pdf",
    "Code of Criminal Procedure 1973": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1973-2.pdf",
    "Indian Evidence Act 1872": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1872-1.pdf",
    "Indian Contract Act 1872": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1872-9.pdf",
    "Transfer of Property Act 1882": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1882-4.pdf",
    "Sale of Goods Act 1930": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1930-3.pdf",
    "Indian Partnership Act 1932": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1932-9.pdf",
    "Companies Act 2013": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A2013-18.pdf",
    "Income Tax Act 1961": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1961-43.pdf",
    "Information Technology Act 2000": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A2000-21.pdf",
    "Consumer Protection Act 2019": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A2019-35.pdf",
    "Right to Information Act 2005": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A2005-22.pdf",
    "Protection of Women from Domestic Violence Act 2005": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A2005-43.pdf",
    "Juvenile Justice (Care and Protection of Children) Act 2015": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A2016-2.pdf",
    "The Rights of Persons with Disabilities Act 2016": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A2016-49.pdf",
    "Indian Succession Act 1925": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1925-39.pdf",
    "Hindu Marriage Act 1955": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1955-25.pdf",
    "Hindu Adoption and Maintenance Act 1956": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1956-78.pdf",
    "Hindu Minority and Guardianship Act 1956": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1956-32.pdf",
    "Hindu Succession Act 1956": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1956-30.pdf",
    "Muslim Women (Protection of Rights on Divorce) Act 1986": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1986-25.pdf",
    "Special Marriage Act 1954": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1954-43.pdf",
    "Indian Divorce Act 1869": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1869-4.pdf",
    "Negotiable Instruments Act 1881": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1881-26.pdf",
    "Limitation Act 1963": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1963-36.pdf",
    "Stamp Act 1899": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1899-2.pdf",
    "Registration Act 1908": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1908-16.pdf",
    "Guardian and Wards Act 1890": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1890-8.pdf",
    "Probation of Offenders Act 1958": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A1958-20.pdf",
    "Bharatiya Nyaya Sanhita 2023": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A2023-1.pdf",
    "Bharatiya Sakshya Adhiniyam 2023": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A2023-2.pdf",
    "Bharatiya Nagarik Suraksha Sanhita 2023": "https://www.indiacode.nic.in/repealedfileopen?rfilename=A2023-3.pdf"
        # Add more acts as needed
    }

    keyword_input = st.text_input(
        "Enter a legal keyword...",
        st.selectbox("Previous Keywords", [""] + st.session_state.search_history)
    ).strip()
    if st.button("🗑️ Clear History"):
        st.session_state.search_history.clear()
        st.success("History cleared")

    if keyword_input:
        detected_acts = [act for act in known_acts if keyword_input.lower() in act.lower() or act.lower() in keyword_input.lower()]
        results_data = []

        if detected_acts:
            st.write(f"### 📂 Search Results for: `{keyword_input}`")
            for act in detected_acts:
                url = known_acts[act]
                st.markdown(f"### {act}")
                st.markdown(f"🔗 [View / Download PDF]({url})")
                pdf_data = fetch_pdf(url, act)
                if pdf_data:
                    st.download_button(f"📄 Download PDF: {act}", pdf_data, act.replace(" ","_")+".pdf","application/pdf")
                results_data.append({"Title": act, "Link": url})

            st.download_button("📄 Export PDF", export_pdf("Indian Legal Acts", results_data), "results.pdf","application/pdf")
            st.download_button("📝 Export DOCX", export_docx("Indian Legal Acts", results_data), "results.docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.download_button("📊 Export Excel", export_excel(pd.DataFrame(results_data)), "results.xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            search_url = f"https://www.indiacode.nic.in/handle/123456789/act-search?query={keyword_input.replace(' ','+')}"
            st.warning(f"⚠️ No known Act detected or IndiaCode may be down. Confirm via Google News or search manually 👉 [Search]({search_url})")

        if keyword_input not in st.session_state.search_history:
            st.session_state.search_history.insert(0, keyword_input)
            st.session_state.search_history = st.session_state.search_history[:10]
    else:
        st.warning("📥 Enter a legal keyword to search.")

# ---------------- FOOTER ----------------
st.markdown("""
<hr style="margin-top:40px;">
<div style='text-align:center; font-size:14px; color:gray;'>
Created by <strong>Naveen Yadav</strong><br>
<em>LLB(Hons) First Year, Sushant University, Gurugram</em><br>
🎓 Student Hub | © 2025
</div>
""", unsafe_allow_html=True)
